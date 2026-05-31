"""DOCX -> pseudo-page text records for unified ingest."""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.table import Table

from app.ingestion.extract import PageRecord, _normalize, doc_id_from_path


logger = logging.getLogger(__name__)

_MIN_SECTION_CHARS = 80
_MAX_SECTION_CHARS = 12000


def _table_to_text(table: Table) -> str:
    lines: list[str] = ["[TABLE]"]
    for row in table.rows:
        cells = [str(cell.text or "").strip() for cell in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))
    lines.append("[/TABLE]")
    return "\n".join(lines)


def _is_heading(paragraph) -> bool:
    style = getattr(paragraph, "style", None)
    if style is None:
        return False
    name = (getattr(style, "name", None) or "").lower()
    return "heading" in name or name.startswith("title")


def _sections_from_docx(path: Path) -> list[str]:
    document = Document(path)
    sections: list[str] = []
    current_title: str | None = None
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines, current_title
        if not current_lines:
            return
        body = "\n".join(current_lines).strip()
        if len(body) < 40:
            current_lines = []
            return
        header = f"# {current_title}\n\n" if current_title else ""
        sections.append(_normalize(header + body))
        current_lines = []

    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        if _is_heading(para):
            flush()
            current_title = text
            continue
        current_lines.append(text)

    for table in document.tables:
        current_lines.append(_table_to_text(table))

    flush()

    if not sections:
        all_text = "\n".join(
            (p.text or "").strip() for p in document.paragraphs if (p.text or "").strip()
        )
        merged = _normalize(all_text)
        if merged:
            sections.append(merged)

    if not sections:
        return []

    merged_sections: list[str] = []
    buffer = ""
    for section in sections:
        if len(buffer) + len(section) <= _MAX_SECTION_CHARS:
            buffer = f"{buffer}\n\n{section}".strip() if buffer else section
            if len(buffer) >= _MIN_SECTION_CHARS:
                merged_sections.append(buffer)
                buffer = ""
            continue
        if buffer:
            merged_sections.append(buffer)
            buffer = section if len(section) <= _MAX_SECTION_CHARS else section[:_MAX_SECTION_CHARS]
        else:
            merged_sections.append(section[:_MAX_SECTION_CHARS])
    if buffer:
        merged_sections.append(buffer)

    return merged_sections


def extract_docx(path: Path) -> Iterator[PageRecord]:
    """Yield one PageRecord per logical section in a Word document."""
    doc_id = doc_id_from_path(path)
    doc_title = path.stem
    sections = _sections_from_docx(path)
    logger.info("[%s] %d DOCX section(s)", path.name, len(sections))

    for index, text in enumerate(sections, start=1):
        if len(text) < 40:
            continue
        yield PageRecord(
            doc_id=doc_id,
            doc_title=doc_title,
            pdf_page=index,
            printed_page=None,
            text=text,
        )


def discover_docx(docs_dir: Path) -> list[Path]:
    files = sorted(
        p for p in docs_dir.glob("*.docx") if p.is_file() and not p.name.startswith("~$")
    )
    logger.info("Found %d DOCX file(s) in %s", len(files), docs_dir)
    return files
